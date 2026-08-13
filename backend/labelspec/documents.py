from __future__ import annotations

import csv
import io
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Literal


SUPPORTED_STANDARD_EXTENSIONS = {".md", ".txt", ".csv", ".docx", ".pdf", ".xlsx"}
MAX_STANDARD_FILES = 20
MAX_STANDARD_FILE_BYTES = 10 * 1024 * 1024
DocumentRole = Literal["auto", "definition", "boundary", "priority"]
DOCUMENT_ROLES = {"auto", "definition", "boundary", "priority"}


@dataclass
class ParsedDocument:
    filename: str
    media_type: str
    raw_content: bytes
    extracted_text: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    role: DocumentRole = "auto"


def infer_document_role(filename: str, text: str = "") -> DocumentRole:
    """Infer the narrowest role we can from a filename/content hint."""
    value = f"{filename}\n{text[:2000]}".lower()
    if any(token in value for token in ("混淆", "边界", "confusion", "boundary")):
        return "boundary"
    if any(token in value for token in ("优先级", "priority")):
        return "priority"
    if any(token in value for token in ("分类标准", "标签定义", "taxonomy", "definition")):
        return "definition"
    return "auto"


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("文本文件编码无法识别，请使用 UTF-8 或 GB18030")


def _parse_pdf(content: bytes) -> tuple[str, Dict[str, Any]]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency installation failure
        raise ValueError("服务端未安装 PDF 解析依赖 pypdf") from exc
    reader = PdfReader(io.BytesIO(content))
    pages: List[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(f"[Page {index}]\n{text}")
    if not pages:
        raise ValueError("PDF 中没有可提取文本；扫描件暂不支持 OCR")
    return "\n\n".join(pages), {"pages": len(reader.pages)}


def _parse_docx(content: bytes) -> tuple[str, Dict[str, Any]]:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise ValueError("服务端未安装 DOCX 解析依赖 python-docx") from exc
    document = Document(io.BytesIO(content))
    blocks: List[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = paragraph.style.name if paragraph.style else ""
        blocks.append(f"[{style}] {text}" if style.startswith("Heading") else text)
    for table_index, table in enumerate(document.tables, start=1):
        blocks.append(f"[Table {table_index}]")
        for row in table.rows:
            blocks.append("\t".join(cell.text.strip() for cell in row.cells))
    if not blocks:
        raise ValueError("DOCX 中没有可提取文本")
    return "\n".join(blocks), {"paragraphs": len(document.paragraphs), "tables": len(document.tables)}


def _parse_xlsx(content: bytes) -> tuple[str, Dict[str, Any]]:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:  # pragma: no cover
        raise ValueError("服务端未安装 XLSX 解析依赖 openpyxl") from exc
    workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    blocks: List[str] = []
    row_count = 0
    for sheet in workbook.worksheets:
        blocks.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value).strip() for value in row]
            if any(values):
                blocks.append("\t".join(values))
                row_count += 1
    if not row_count:
        raise ValueError("XLSX 中没有可提取内容")
    return "\n".join(blocks), {"sheets": len(workbook.worksheets), "rows": row_count}


def parse_standard_document(
    filename: str, media_type: str, content: bytes, role: DocumentRole = "auto"
) -> ParsedDocument:
    if role not in DOCUMENT_ROLES:
        raise ValueError(f"无效的标准文档角色: {role}")
    safe_name = Path(filename).name
    suffix = Path(safe_name).suffix.lower()
    if suffix not in SUPPORTED_STANDARD_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_STANDARD_EXTENSIONS))
        raise ValueError(f"不支持 {suffix or '无扩展名'} 标准文档，仅支持: {supported}")
    if not content:
        raise ValueError(f"{safe_name} 是空文件")
    if len(content) > MAX_STANDARD_FILE_BYTES:
        raise ValueError(f"{safe_name} 超过 10 MB 限制")

    metadata: Dict[str, Any] = {"size": len(content), "extension": suffix}
    if suffix == ".pdf":
        text, parsed_metadata = _parse_pdf(content)
    elif suffix == ".docx":
        text, parsed_metadata = _parse_docx(content)
    elif suffix == ".xlsx":
        text, parsed_metadata = _parse_xlsx(content)
    else:
        text = _decode_text(content).strip()
        parsed_metadata = {}
        if suffix == ".csv":
            rows = list(csv.reader(io.StringIO(text)))
            parsed_metadata = {"rows": len(rows)}
    if not text.strip():
        raise ValueError(f"{safe_name} 中没有有效文本")
    metadata.update(parsed_metadata)
    effective_role = infer_document_role(safe_name, text) if role == "auto" else role
    return ParsedDocument(
        filename=safe_name,
        media_type=media_type or "application/octet-stream",
        raw_content=content,
        extracted_text=text.strip(),
        metadata=metadata,
        role=effective_role,
    )
