from io import BytesIO

import pytest
from docx import Document
from openpyxl import Workbook

from labelspec.documents import infer_document_role, parse_standard_document
from labelspec.api import _parse_upload, dataset_template


def test_text_and_office_documents_are_extracted() -> None:
    markdown = parse_standard_document("rules.md", "text/markdown", "# 标签\n贷款定义".encode())
    assert "贷款定义" in markdown.extracted_text

    document = Document()
    document.add_heading("汽车", level=1)
    document.add_paragraph("购车定义")
    docx = BytesIO()
    document.save(docx)
    parsed_docx = parse_standard_document(
        "rules.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", docx.getvalue()
    )
    assert "购车定义" in parsed_docx.extracted_text

    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["一级", "二级", "定义"])
    sheet.append(["金融", "贷款", "借贷相关"])
    xlsx = BytesIO()
    workbook.save(xlsx)
    parsed_xlsx = parse_standard_document(
        "taxonomy.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", xlsx.getvalue()
    )
    assert "金融\t贷款\t借贷相关" in parsed_xlsx.extracted_text


def test_unsupported_document_is_rejected() -> None:
    with pytest.raises(ValueError, match="不支持"):
        parse_standard_document("rules.exe", "application/octet-stream", b"content")


def test_document_role_is_inferred_from_filename() -> None:
    assert infer_document_role("example-confusion-boundary.md") == "boundary"
    assert infer_document_role("priority-rules.md") == "priority"
    assert infer_document_role("分类标准.md") == "definition"
    assert infer_document_role("notes.md") == "auto"


def test_dataset_upload_formats() -> None:
    assert _parse_upload("cases.txt", "第一条\n\n第二条".encode()) == [{"text": "第一条"}, {"text": "第二条"}]
    assert _parse_upload("cases.csv", "text,gold_label\n第一条,标签A\n".encode()) == [{"text": "第一条", "gold_label": "标签A"}]
    assert _parse_upload("cases.jsonl", '{"text":"第一条"}\n'.encode()) == [{"text": "第一条"}]
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["text", "gold_label"])
    sheet.append(["第二条", "标签B"])
    xlsx = BytesIO()
    workbook.save(xlsx)
    assert _parse_upload("cases.xlsx", xlsx.getvalue()) == [{"text": "第二条", "gold_label": "标签B"}]


def test_dataset_upload_requires_text_field() -> None:
    with pytest.raises(ValueError, match="text 表头"):
        _parse_upload("cases.csv", "content\n第一条\n".encode())
    with pytest.raises(ValueError, match="缺少 text"):
        _parse_upload("cases.jsonl", '{"content":"第一条"}\n'.encode())


def test_dataset_template_contains_optional_gold_label() -> None:
    response = dataset_template()
    assert response.status_code == 200
    assert response.body.decode().startswith("text,gold_label")
